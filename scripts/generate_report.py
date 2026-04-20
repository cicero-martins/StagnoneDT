"""Generate the Stagnone DT progress report (Word) from project sources.

Consolidates:
- memory files (project context, lessons, feedback)
- notebook titles and first markdown cells
- existing README.md and docs/EDITO_WORKFLOW.md
- key figures from figures/

Output: f:/StagnoneDT/docs/progress_report_v1.docx

Run:
    python scripts/generate_report.py
"""
from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = PROJECT_ROOT / 'docs'
NOTEBOOKS_DIR = PROJECT_ROOT / 'notebooks'
FIGURES_DIR = PROJECT_ROOT / 'figures'
MEMORY_DIR = Path(r'C:\Users\cicer\.claude\projects\f--StagnoneDT\memory')
OUTPUT = DOCS_DIR / 'progress_report_v1.docx'


# ---------- helpers ----------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def add_figure(doc, path: Path, caption: str, width_cm: float = 14.0):
    if not path.exists():
        add_para(doc, f'[Figure not found: {path.name} — see {path}]', italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    try:
        r.add_picture(str(path), width=Cm(width_cm))
    except Exception as e:
        add_para(doc, f'[Figure embedding failed for {path.name}: {e}]', italic=True)
        return
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f'Figure: {caption}')
    cr.italic = True
    cr.font.size = Pt(10)


def add_page_break(doc):
    doc.add_page_break()


def read_memory(filename: str) -> tuple[dict, str]:
    """Read a memory markdown file, split frontmatter + body."""
    path = MEMORY_DIR / filename
    if not path.exists():
        return {}, ''
    text = path.read_text(encoding='utf-8')
    if not text.startswith('---'):
        return {}, text
    parts = text.split('---', 2)
    if len(parts) < 3:
        return {}, text
    fm_raw = parts[1].strip()
    body = parts[2].strip()
    fm = {}
    for line in fm_raw.splitlines():
        if ':' in line:
            k, v = line.split(':', 1)
            fm[k.strip()] = v.strip()
    return fm, body


def read_notebook_first_markdown(nb_path: Path) -> tuple[str, str, int]:
    """Return (title, purpose_lines, cell_count) from first markdown cell."""
    try:
        nb = json.loads(nb_path.read_text(encoding='utf-8'))
    except Exception:
        return (nb_path.stem, '(unreadable)', 0)
    cells = nb.get('cells', [])
    n_cells = len(cells)
    for c in cells:
        if c.get('cell_type') == 'markdown':
            src = c.get('source', '')
            if isinstance(src, list):
                src = ''.join(src)
            lines = [l.strip() for l in src.splitlines() if l.strip()]
            if not lines:
                continue
            # first line = title (strip leading #)
            title = lines[0].lstrip('#').strip()
            purpose = ' '.join(lines[1:4]) if len(lines) > 1 else ''
            return (title, purpose, n_cells)
    return (nb_path.stem, '', n_cells)


def set_cell_bg(cell, hex_color: str):
    """Shade a table cell. hex_color without '#'."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_table(doc, header: list[str], rows: list[list[str]], col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        set_cell_bg(cell, 'D9E1F2')
    # rows
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = str(val)
    # col widths
    if col_widths_cm:
        for row in t.rows:
            for j, w in enumerate(col_widths_cm):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)
    return t


# ---------- report sections ----------

def write_cover(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Stagnone di Marsala — Digital Twin')
    r.bold = True
    r.font.size = Pt(24)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Internal technical progress report — v1')
    r.italic = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run(f'Compiled on: {date.today().isoformat()}').font.size = Pt(11)

    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth.add_run('Cicero Martins Jr — Università degli Studi di Palermo').font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        'Scope: 3D coupled wave-hydrodynamic modeling (D-Flow FM + SWAN) of a shallow '
        'hypersaline Mediterranean lagoon, deployed as a digital twin on the EDITO Datalab.'
    ).font.size = Pt(11)

    add_page_break(doc)


def write_toc_placeholder(doc):
    add_heading(doc, 'Table of contents', level=1)
    add_para(doc,
             '[In Word: Insert → Table of Contents → Automatic Table 1. '
             'Right-click → Update Field after editing.]',
             italic=True)
    add_page_break(doc)


def write_executive_summary(doc):
    add_heading(doc, '1. Executive summary', level=1)

    add_para(doc,
             'This report documents ~2 weeks of work building a 3D coupled wave-hydrodynamic '
             'digital twin of the Stagnone di Marsala lagoon (Sicily, Italy), targeting the '
             'EDITO Datalab for production runs.')

    add_heading(doc, 'Problem', level=2)
    add_para(doc,
             'The Stagnone di Marsala is a shallow (0.5–2 m) hypersaline (~42 psu) Mediterranean '
             'lagoon supporting Posidonia/Cymodocea seagrass meadows. Despite microtidal forcing, '
             'the lagoon exhibits significant vertical flow structure driven by wind — a feature '
             'that demands 3D modeling. Previous work had a modelbuilder-generated configuration '
             'that did not run; open questions about residence time, inlet exchange, seagrass '
             'drag, and turbidity plumes from adjacent canals motivate this digital twin.')

    add_heading(doc, 'Approach', level=2)
    add_para(doc,
             'Starting from the Model B reference (dfm_tools modelbuilder), we progressively '
             'built three model versions: v01 (baseline, fixed to run), v02 (calibration corrections '
             'for datum, drying, wind), v03 (waves, residence-time tracer, hypersaline initial '
             'condition). Work is organized across 13 operational Jupyter notebooks covering '
             'mesh evaluation, in-situ QC, satellite roughness, model build, calibration, '
             'diagnostics, particle tracking, residence time, and EDITO post-processing.')

    add_heading(doc, 'Current status', level=2)
    add_para(doc,
             'v01 complete locally. v02 running on EDITO (after a local power-outage incident '
             'prompted migration). v03 built and pending v02 validation. EDITO workflow '
             'operational: S3 upload of inputs (~140 MB), container-based execution via '
             'delft3dfm_run_docker (2D3D-HMWQ variant), streaming post-processing via s3fs '
             '(no 16 GB map.nc download required).')

    add_heading(doc, 'Next steps', level=2)
    add_para(doc,
             'Validate v02 outputs; build v03 and run short sanity check; download extended '
             'CMEMS/ERA5 for the ≥30-day residence-time run; verify pod CPU allocation for '
             'parallel (nPart=4) execution on EDITO; run particle tracking against the July 2025 '
             'drifter campaign; tune wind drag if surface drift is underpredicted.')

    add_page_break(doc)


def write_introduction(doc):
    add_heading(doc, '2. Introduction', level=1)

    # project memory
    _, project_body = read_memory('project_stagnone_dt.md')

    add_heading(doc, '2.1 Site description', level=2)
    add_para(doc,
             'The Stagnone di Marsala is a shallow coastal lagoon on Sicily\'s western coast '
             '(center ~37.86°N, 12.46°E), with an area of ~12 km² and water depths of 0.5–2 m. '
             'The tidal range is microtidal (~0.2–0.3 m). Two inlets connect the lagoon to the '
             'open Mediterranean: Boca Nord (shallow, ~0.3 m) and Boca Sud (deeper, ~1.5 m). '
             'The lagoon is a protected area with extensive seagrass cover (Posidonia oceanica '
             'and Cymodocea nodosa) and is characterized as hypersaline: interior salinity '
             '~42 psu vs offshore Mediterranean ~37.5 psu. Despite its shallow depth, significant '
             'vertical flow structure is observed — wind-driven surface currents contrasting with '
             'bottom currents (evidenced by plant orientation).')

    add_heading(doc, '2.2 Research questions', level=2)
    add_para(doc, 'The digital twin is designed to address five interconnected questions:')
    for q in [
        'What is the 3D vertical flow structure of the lagoon, and how is it driven by wind forcing?',
        'What is the characteristic residence time of lagoon water, and how does it vary with wind regimes?',
        'How do the two inlets exchange water with the open Mediterranean, and what is the role of the hypersaline gradient?',
        'How do seagrass beds modify bottom friction and wave attenuation, and what does this imply for sediment transport?',
        'How do episodic turbidity plumes (from the Trapani salt pans and airport canals) propagate through the lagoon?',
    ]:
        doc.add_paragraph(q, style='List Bullet')

    add_heading(doc, '2.3 Planned publications', level=2)
    add_para(doc, 'Five Q1 journal papers are planned, organized around the research questions:')
    papers = [
        ('Paper 1', '3D hydrodynamic modeling of a shallow Mediterranean seagrass lagoon: vertical flow structure under wind forcing', 'Estuarine, Coastal and Shelf Science / Continental Shelf Research'),
        ('Paper 2', 'Satellite-derived bottom roughness for shallow lagoon modeling: seagrass classification to Manning coefficients', 'Remote Sensing of Environment / Ecological Modelling'),
        ('Paper 3', 'Coupled wave-hydrodynamic modeling in a seagrass lagoon: wave-vegetation interaction and turbidity implications', 'Ocean Modelling / Coastal Engineering'),
        ('Paper 4', 'Multi-source bathymetry fusion for vegetated shallow lagoons', 'Remote Sensing / Journal of Coastal Research'),
        ('Paper 5', 'A digital twin framework for protected Mediterranean lagoons on the EDITO platform', 'Environmental Modelling & Software / Science of the Total Environment'),
    ]
    add_table(doc,
              header=['#', 'Title', 'Target journals'],
              rows=[[p[0], p[1], p[2]] for p in papers],
              col_widths_cm=[1.5, 10, 5])

    add_page_break(doc)


def write_methodology(doc):
    add_heading(doc, '3. Methodology', level=1)

    add_heading(doc, '3.1 Architecture', level=2)
    add_para(doc,
             'The digital twin couples a 3D D-Flow FM (flexible-mesh, unstructured) hydrodynamic '
             'core with a SWAN spectral wave model. Components communicate via DIMR (Delft '
             'Integrated Model Runner). Forcing and boundary data come from CMEMS (ocean), ERA5 '
             '(atmosphere), TPXO8 (tides), and in-situ stations (3 water-level + 2 wind). '
             'Validation uses the same in-situ network plus July 2025 Stokes drifter deploys '
             'and Sentinel-2 satellite imagery.')

    add_heading(doc, '3.2 Reference models', level=2)
    add_para(doc,
             'Two pre-existing model configurations served as starting points:')
    add_para(doc, 'Model A (simple, oldModel/input/):', bold=True)
    add_para(doc,
             'Lagoon-only model, boundaries at the inlets driven by in-situ water levels. CRS '
             'UTM Zone 32 (EPSG:32632). 5 sigma layers, k-epsilon, no salinity/temperature, no '
             'waves. 8 polygon friction zones (Manning n=0.05). Status: was running; limitation '
             'is reliance on inlet measurements rather than offshore forcing.')
    add_para(doc, 'Model B (modelbuilder-based, oldModel/Stagnone_py_lagoon3D.dsproj_data/):', bold=True)
    add_para(doc,
             'Generated with dfm_tools modelbuilder, offshore boundary ~10-15 km west, 49 '
             'boundary points. WGS84 (EPSG:4326), 10 sigma layers, k-epsilon, salinity + '
             'temperature enabled, CMEMS boundary conditions (WL + tidal TPXO8 + S/T + '
             'velocity), ERA5 meteo. 19 observation points. Period: July 1–10, 2025. '
             'Status: not running — priority target for v01.')

    add_heading(doc, '3.3 Model version comparison', level=2)
    add_table(doc,
              header=['Feature', 'v01', 'v02', 'v03', 'v03a'],
              rows=[
                  ['Datum offset (WL BC)', '0.0 m', '+0.42 m', '+0.42 m', '+0.42 m'],
                  ['Obs points BS / AE', 'drying', 'always-wet', 'always-wet', 'always-wet'],
                  ['Central 3D obs', 'none', 'C1, C2, C3', 'C1, C2, C3', 'C1, C2, C3'],
                  ['Wind forcing', 'ERA5 only', 'blended', 'blended', 'blended'],
                  ['Initial salinity', 'uniform 37.5', 'uniform 37.5', 'attempted 42 interior (not applied — nudging)', 'attempted 42 interior (still not applied)'],
                  ['Passive tracer', 'no', 'no', 'yes', 'yes'],
                  ['SWAN wave coupling', 'no', 'no', 'intended (constant BC, single grid)', 'yes, nested: outer 800 m + inner 100 m'],
                  ['Simulation period', 'Jul 1–10', '9 d', '9 d', '9 d'],
                  ['Execution', 'local, seq.', 'EDITO, seq.', 'EDITO, 4× MPI', 'local, 4× MPI'],
              ],
              col_widths_cm=[3, 3, 3, 3, 3])

    add_heading(doc, '3.4 Data inputs', level=2)
    inputs = [
        ('Mesh', 'Stagnone_dxy01_15m_net.nc from Model B, 25,358 faces, WGS84, hand-refined lagoon margins'),
        ('Bathymetry', '2006 XYZ survey (189K points, 20 m grid, Sicily coast) + GEBCO offshore'),
        ('In-situ WL', 'BocaNord (37.905°N, 12.457°E), BocaSud (37.847°N, 12.449°E), AltaVilaEst (37.890°N, 12.447°E); 10-min UTC'),
        ('In-situ wind', 'AltaVilaEst + Mulino stations; 3 m anemometer height, converted to 10 m via log profile (factor ~1.18)'),
        ('Drifters', '12 Stokes drifter deploys, July 8–9 2025, surface positions in UTM Zone 32N'),
        ('CMEMS', 'MEDSEA_MULTIYEAR_PHY_006_004: WL, T, S, velocity at boundary'),
        ('ERA5', 'u10n, v10n, msl, chnk (Charnock parameter)'),
        ('Tides', 'TPXO8 harmonics at boundary'),
        ('Satellite', 'Sentinel-2 L2A for seagrass classification (notebook 03)'),
    ]
    add_table(doc,
              header=['Input', 'Source / properties'],
              rows=inputs,
              col_widths_cm=[4, 13])

    # Figure: mesh / domain overview
    add_figure(doc, FIGURES_DIR / 'domain_overview.png',
               'Model domain and mesh overview (v01, 25k faces).')

    add_page_break(doc)


def write_implementation(doc):
    add_heading(doc, '4. Implementation', level=1)

    # 4.1 Notebook catalog
    add_heading(doc, '4.1 Notebook catalog', level=2)
    add_para(doc, 'The project workflow is organized in 13 Jupyter notebooks:')

    nb_table = []
    for nb_file in sorted(NOTEBOOKS_DIR.glob('*.ipynb')):
        title, _, n_cells = read_notebook_first_markdown(nb_file)
        nb_table.append([nb_file.name, title, str(n_cells)])
    add_table(doc,
              header=['Notebook', 'Title / purpose', 'Cells'],
              rows=nb_table,
              col_widths_cm=[5, 11, 1.5])

    # 4.2 Model B fixes
    add_heading(doc, '4.2 Model B fixes — getting v01 to run', level=2)
    add_para(doc,
             'The modelbuilder-generated Model B did not run out of the box. Five issues were '
             'identified and fixed in notebook 04 to produce a running v01:')
    fixes = [
        ('Missing output file names', 'HisFile and MapFile were blank', 'Specified output NetCDF filenames'),
        ('Background vertical viscosity too small', 'Vicoww = Dicoww = 5e-5', 'Raised to 1e-4 for numerical stability in shallow cells'),
        ('Initial temperature wrong', '6 °C (wrong for July Mediterranean)', 'Set to 24 °C'),
        ('Initial salinity low', '30 psu', 'Set to 37.5 psu (offshore Med summer)'),
        ('Unsupported keywords', 'Umodlin, EffectSpiral, WaveNikuradse, S1incinterval', 'Stripped from MDU before hydrolib-core load'),
    ]
    add_table(doc,
              header=['Issue', 'Original value', 'Fix'],
              rows=fixes,
              col_widths_cm=[6, 5, 6])

    # 4.3 v01 diagnostics
    add_heading(doc, '4.3 v01 diagnostics — four calibration findings', level=2)
    add_para(doc,
             'After v01 completed its 9-day run, notebooks 08–09 produced diagnostics that '
             'revealed four systematic issues driving v02 changes:')
    for i, line in enumerate([
        '1. Systematic negative datum bias: the model under-predicted water level at all three stations by ~0.4 m, consistent across the run. Root cause: the constant WL boundary had no offset to align the model datum to the observed MSL reference.',
        '2. Amplitude flattening at BocaSud and AltaVilaEst: the model time series at these stations showed damped tidal amplitude compared to observations. Root cause: both obs points fell in cells that dried intermittently, so the model time series was contaminated by drying artifacts.',
        '3. ERA5 wind underpredicted at lagoon scale: ERA5 wind magnitude was systematically lower than in-situ observations at AE and Mulino. Root cause: ERA5 grid (~25 km) does not resolve coastal/topographic effects.',
        '4. Missing central observation points: the 19 obs points in Model B did not cover the lagoon interior adequately to evaluate 3D vertical structure — the wind-driven surface/bottom contrast that motivates the 3D approach.',
    ]):
        doc.add_paragraph(line)

    add_figure(doc, FIGURES_DIR / 'insitu_wl_comparison.png',
               'v01 water-level comparison at BN/BS/AE stations, showing the ~0.4 m datum bias '
               'and amplitude damping.')
    add_figure(doc, FIGURES_DIR / 'diag_bias_correction.png',
               'Datum-offset correction test (notebook 09): applying the mean bias as a constant '
               'offset removes the systematic under-prediction.')

    # 4.4 v02 changes
    add_heading(doc, '4.4 v02 changes — four fixes', level=2)
    changes = [
        ('Datum offset', 'Set constant WL boundary to +0.4208 m (mean v01 bias across BN/BS/AE)'),
        ('Replacement obs cells', 'Moved BocaSud and AltaVilaEst obs to deepest always-wet cells within 500 m (min depth > 0.15 m)'),
        ('Central 3D obs', 'Added C1 (12.455, 37.870), C2 (12.460, 37.890, user-adjusted), C3 (12.458, 37.855)'),
        ('Blended wind', 'ERA5 outside lagoon + IDW of AE + Mulino inside; inner radius 3 km, outer 8 km, 500 m output grid'),
    ]
    add_table(doc,
              header=['Change', 'Implementation'],
              rows=changes,
              col_widths_cm=[5, 12])
    add_figure(doc, FIGURES_DIR / 'wind_blend_snapshot.png',
               'Snapshot of blended wind field (notebook 11): ERA5 offshore + in-situ stations '
               'inside the lagoon via IDW, with smooth transition band.')
    add_figure(doc, FIGURES_DIR / 'diag_replacement_cells.png',
               'Replacement cells for BS/AE obs points (notebook 10): dry-cell analysis identifies '
               'deepest always-wet neighbors within 500 m.')

    # 4.5 v03 design
    add_heading(doc, '4.5 v03 design', level=2)
    add_para(doc,
             'v03 layers four additional corrections on v02. Built end-to-end in notebook 14:')
    for c in [
        'Hypersaline initial salinity: 42 psu inside the lagoon polygon, 37.5 psu outside. The polygon was extracted authoritatively from the oldModel/Stagnone_justLagoon mesh (hand-drawn), reprojected from EPSG:3857 to WGS84 — 82 simplified vertices.',
        'Passive residence-time tracer "lagoon_tracer": initial value 1.0 inside lagoon / 0.0 outside, Dirichlet=0 at the open ocean boundary. Decay of the spatial mean yields the e-folding residence time.',
        'SWAN wave coupling activated (Wavemodelnr=3). Notebook 05 supplies the mdw, grid, and dimr_config with both FM and Wave components.',
        'Extended simulation period: requires downloading CMEMS + ERA5 for at least 30 days to capture full tracer decay. Currently blocked pending data download.',
    ]:
        doc.add_paragraph(c, style='List Bullet')

    # 4.6 EDITO deployment
    add_heading(doc, '4.6 EDITO deployment', level=2)
    add_para(doc,
             'The EDITO Datalab hosts the delft3dfm_run_docker process (variant "2D3D-HMWQ") '
             'which accepts a model folder at s3://<user-bucket>/DFM_INPUT/ and writes outputs '
             'to s3://<user-bucket>/DFM_OUTPUT/. User quota: 32 GiB RAM, 8 CPU, 50 GiB storage.')
    add_para(doc,
             'A Python sync script (scripts/edito_sync.py) handles: listing buckets, cleaning '
             'old outputs, uploading model inputs (~140 MB), uploading project code, downloading '
             'results. Credentials come from .env (gitignored) with values from the EDITO '
             'storage settings page.')
    add_para(doc, 'Critical gotcha (documented in memory/feedback_edito_run_model_sh.md):',
             bold=True)
    add_para(doc,
             'Every uploaded model must include run_model.sh at its root with LF line endings. '
             'Missing it produces a misleading Kubernetes error ("MountVolume.SetUp failed ... '
             'kube-root-ca.crt not registered") that looks like an infrastructure problem but '
             'is actually the container failing to find its entrypoint. The v02 upload initially '
             'failed here; after adding the dfm_tools modelbuilder-style shell script, the run '
             'started immediately.')
    add_para(doc,
             'Post-processing on EDITO uses JupyterLab from the service catalog. Notebook 15 '
             'demonstrates streaming the 16 GB map.nc via s3fs + xarray.open_dataset with '
             'chunked reads — critical to avoid downloading large files within quota.')

    add_page_break(doc)


def write_results(doc):
    add_heading(doc, '5. Preliminary results', level=1)

    add_heading(doc, '5.1 v01 validation', level=2)
    add_para(doc,
             'v01 completed its 9-day simulation (July 1–10, 2025) locally in ~4 hours of '
             'wall-clock time. Water-level comparison at the three in-situ stations (BN/BS/AE) '
             'revealed the systematic bias and amplitude issues addressed in v02.')
    add_figure(doc, FIGURES_DIR / 'model_vs_obs_timeseries.png',
               'v01 time series vs observations at BN/BS/AE (notebook 08). The ~0.4 m offset '
               'and amplitude damping at BS/AE drove the v02 fixes.')
    add_figure(doc, FIGURES_DIR / 'model_vs_obs_scatter.png',
               'v01 scatter comparison at the three stations. Points cluster below the 1:1 line '
               'consistent with negative bias.')

    add_heading(doc, '5.2 v02 partial results (local)', level=2)
    add_para(doc,
             'v02 completed approximately 70% of its 9-day run locally before a power outage '
             'interrupted execution (stop at simulated day 6, ~July 7 08:00 UTC). The outputs '
             'through that point are intact: 912 time steps in his.nc (10-min interval), 305 '
             'snapshots in map.nc (30-min interval). The final map.nc snapshot has a corrupted '
             'timestamp from the partial write — dropped in post-processing.')
    add_para(doc,
             'This partial run is usable for WL validation and v01 comparison, but does not '
             'cover the July 8–9 drifter campaign window needed for particle-tracking '
             'validation.')

    add_heading(doc, '5.3 v02 on EDITO', level=2)
    add_para(doc,
             'After the power-outage incident, v02 was migrated to EDITO. Upload path: 140 MB '
             'of inputs to s3://oidc-cmartinsjr/DFM_INPUT/ via boto3. v02 completed the 9-day '
             'run on EDITO and produced dramatically improved WL metrics (RMSE 6-7× lower than '
             'v01, bias near zero, Willmott ~0.85 at BocaNord). The +0.42 m datum offset fixed '
             'the systematic bias; BocaSud still shows residual RMSE ~0.10 m and Willmott 0.66 '
             '(under investigation).')

    add_heading(doc, '5.4 v03 on EDITO (coupled waves + tracer)', level=2)
    add_para(doc,
             'v03 added three upgrades over v02: passive residence-time tracer, attempted '
             'hypersaline salinity initial field (42 psu interior), and intended SWAN wave '
             'coupling. Ran on EDITO with nPart=4 MPI processes; completed the 9-day window '
             'in ~5 hours wall-clock. Outputs: partitioned his.nc + 4× map.nc (~9.6 GB each).')
    add_para(doc,
             'Findings: (a) WL metrics unchanged vs v02 — waves had negligible impact on '
             'water level in this microtidal regime; (b) hypersaline initial field NOT applied — '
             'salinity stayed at ~37.7 psu interior (the nudging file or CMEMS salinitybnd '
             'overrides the XYZ initial field); (c) tracer decay successful: e-folding time '
             'τ ≈ 3.9 d at C1_Central (interior), 4.6 d at C3_SouthCenter, <1 d near inlets. '
             'This is the first concrete estimate of lagoon residence time.')

    add_heading(doc, '5.5 v03a local (nested SWAN active + 4× MPI)', level=2)
    add_para(doc,
             'v03a is the authoritative version: inputs were updated on the office computer '
             '(inner SWAN grid regenerated at 121×221 @ 100 m, simulation bumped to full 9 days '
             'with daily restart files), then locally we added the outer SWAN grid at 69×54 '
             '@ ~800 m covering the entire FM domain, with NestedInDomain=1 on the inner — '
             'now the full FM domain has wave coupling instead of just the lagoon.')
    add_para(doc,
             'Two non-obvious fixes were required for SWAN-WAVE to accept the grids: (a) the '
             '.dep files must use FORTRAN fixed-width format "12 values per line, %17.7e", '
             'exactly matching the inner file that worked; any extra whitespace triggers a '
             'misleading "Premature end of file" error; (b) both depth files must be padded '
             'to (MC+1)×(NC+1) with -99 fill values in the extra row and column (end-of-row/'
             'column markers required by the Delft3D-WAVE parser).')
    add_para(doc,
             'The run completed 9 days locally in 5h38min on a Ryzen 5 3600 (6 physical / '
             '12 logical cores, 16 GB RAM) with 4 MPI processes. Total output: 27 GB across '
             'four map.nc partitions plus SWAN output on both grids.')
    add_para(doc, 'Results:', bold=True)
    add_para(doc,
             '(1) WL metrics essentially identical to v02 / v03 (±0.002 m RMSE, ±0.01 Willmott), '
             'confirming the waves contribute negligibly to the water level signal. (2) '
             'Hypersaline salinity STILL not applied — same 37.7 psu throughout the lagoon. '
             'The fix of setting initialSalinity=0.0 in the MDU did not help, suggesting the '
             'nudging file or boundary condition is overriding. Requires investigation of the '
             'iniWithNudge flag and the nudge_salinity_temperature initial condition file. '
             '(3) Tracer decay consistent with v03 EDITO: τ ≈ 3.0 d at C1_Central and 3.7 d at '
             'C3_SouthCenter — residence time estimate robust across runs.')
    add_para(doc, 'Key achievement:', bold=True)
    add_para(doc,
             'The model runs end-to-end locally with FM + nested SWAN coupling via DIMR, '
             'producing consistent validation metrics and usable tracer time series. This '
             'establishes a reliable baseline for subsequent scenario runs (canal discharge, '
             'wind ensemble, seagrass roughness) without relying on EDITO compute availability.')
    add_figure(doc, FIGURES_DIR / 'v03a_wl_validation.png',
               'v03a water-level validation at BN/BS/AE. Metrics essentially identical to v02 '
               'and v03.')
    add_figure(doc, FIGURES_DIR / 'v03a_salinity_check.png',
               'v03a surface salinity evolution — the hypersaline initial condition (target '
               '42 psu) is not being applied; interior stays at ~37.7 psu throughout. Red '
               'dotted line is the intended interior value.')

    add_heading(doc, '5.6 Wind blending', level=2)
    add_para(doc,
             'The ERA5+in-situ wind blend (notebook 11) produces a composite field used from '
             'v02 onward. Inside a 3 km radius of each in-situ station, the IDW interpolation '
             'of station observations dominates; between 3 and 8 km, a smooth weight transitions '
             'back to ERA5; beyond 8 km, ERA5 is used directly. Grid resolution of the blended '
             'product: 500 m.')
    add_figure(doc, FIGURES_DIR / 'wind_era5_vs_insitu.png',
               'Comparison of ERA5 and in-situ wind at AE station (notebook 11). ERA5 systematically '
               'underpredicts wind speed inside the lagoon at the ~25 km grid scale.')

    add_page_break(doc)


def write_extensions(doc):
    add_heading(doc, '6. Research extensions', level=1)

    add_heading(doc, '6.1 Particle tracking — July 2025 drifter campaign', level=2)
    add_para(doc,
             'In July 2025, 12 Stokes drifter deploys were performed in the lagoon. Positions '
             'were recorded every few minutes in UTM Zone 32N. Notebook 12 builds the pipeline '
             'to reproduce these deploys via Lagrangian particle tracking driven by the v02 (or '
             'later) velocity field.')
    add_para(doc,
             'The chosen execution path is OpenDrift offline, reading v02 map.nc and seeding '
             'particles at each deploy\'s first valid observation. Skill metrics per drifter: '
             'haversine endpoint separation, path-length ratio, mean heading bias, and Liu & '
             'Weisberg (2011) trajectory skill score. D-Flow FM\'s built-in particle module is '
             'a fallback if surface drift requires same-numerics consistency.')
    add_para(doc,
             'Notebook 12 execution is gated on v02 validation — the current v02 partial output '
             'stops before the July 8–9 campaign window.')

    add_heading(doc, '6.2 Residence time', level=2)
    add_para(doc,
             'A core research question is the lagoon\'s water residence time. Notebook 13 '
             'implements the primary method — an Eulerian passive tracer initialized to 1 '
             'inside the lagoon, 0 outside, with Dirichlet=0 at the open boundary. The spatial '
             'mean inside the lagoon decays exponentially; the e-folding time is the bulk '
             'residence timescale. Per-cell fits produce spatial maps of local residence time.')
    add_para(doc,
             'Literature review supports this as the standard approach for wind-dominated '
             'lagoons (García-Oliva 2019 for Mar Menor; Cucco & Umgiesser 2006 for Venice; '
             'Ferrarin 2013 for Sacca di Goro). A Knudsen salt-balance estimate τ = V·ΔS / (E·S) '
             'serves as an independent bulk sanity check for the hypersaline case. The gold-'
             'standard CART age tracer (Deleersnijder/Delhez) is a future extension requiring '
             'two coupled user tracers.')
    add_para(doc,
             'Execution requires extending the simulation to ≥30 days (otherwise the tracer '
             'barely starts decaying). This is gated on downloading extended CMEMS + ERA5 data.')

    add_heading(doc, '6.3 Feedback loop — hypersaline correction', level=2)
    add_para(doc,
             'An important diagnostic finding drove v03 design: v01 and v02 both had uniform '
             'initial salinity of 37.5 psu and CMEMS boundary at similar values — meaning the '
             'lagoon interior was modeled as hyposaline, the opposite of the observed hypersaline '
             'state (~42 psu). This has three consequences that v03 must address:')
    for c in [
        'Density-driven exchange at the inlets is inverted in v01/v02 — bottom outflow (denser hypersaline water leaving) is missing.',
        'Salinity cannot serve as a natural diagnostic tracer for residence time until the gradient is correctly represented.',
        'Any coupled biogeochemistry or seagrass dynamics downstream start from a wrong baseline.',
    ]:
        doc.add_paragraph(c, style='List Bullet')
    add_para(doc,
             'The v03 fix (notebook 14) applies a polygon-masked initial field: 42 psu inside '
             'the authoritative lagoon polygon (extracted from oldModel/Stagnone_justLagoon), '
             '37.5 psu outside. An ERA5 evaporation forcing term is pending to close the '
             'steady-state salt balance.')

    add_page_break(doc)


def write_status(doc):
    add_heading(doc, '7. Current status and next steps', level=1)

    add_heading(doc, '7.1 Status snapshot', level=2)
    rows = [
        ['v01 baseline', 'Complete', '9-day run, 4 h wall-clock locally'],
        ['v02 (local)', 'Partial', 'Power outage at 70% (6.3 d simulated intact)'],
        ['v02 (EDITO)', 'Complete', 'Full 9-d; RMSE 6-7× better than v01'],
        ['v03 (EDITO)', 'Complete', 'FM+SWAN intended, tracer working, salinity not applied'],
        ['v03a (local)', 'Complete', '9-d run in 5h38min on 4 MPI; nested SWAN active'],
        ['EDITO workflow', 'Operational', 'S3 sync + JupyterLab + MinIO compat fixes'],
        ['Local parallel run', 'Operational', '4× MPI on Ryzen 5 3600 + SWAN wave.dll'],
        ['Residence time estimate', 'First result', 'τ ≈ 3-4 d interior; consistent v03/v03a'],
        ['Hypersaline salinity', 'Not applied', 'Nudging file overrides XYZ — v04 fix needed'],
        ['Notebook 12 (particles)', 'Framework ready', 'Can now run on v03 or v03a map.nc'],
    ]
    add_table(doc,
              header=['Component', 'Status', 'Notes'],
              rows=rows,
              col_widths_cm=[5, 5, 7])

    add_heading(doc, '7.2 Immediate next steps', level=2)
    for step in [
        'Debug hypersaline initial condition: check iniWithNudge flag in the MDU and the nudge_salinity_temperature file — build v04 where the XYZ override actually takes effect.',
        'Analyze v03a SWAN output: read wavm-stagnone-swan_grid.nc (inner) and wavm-stagnone-swan_grid_outer.nc (outer) to check Hs, Tp, direction across the domain.',
        'Run particle-tracking (notebook 12) against v03 or v03a velocities, covering the July 8-9 drifter campaign window.',
        'Download extended CMEMS + ERA5 (Jul 1 – Aug 15 2025) for a ≥30-day residence-time run to let the tracer decay further.',
        'Upgrade SWAN boundary conditions from constant parametric to CMEMS MEDSEA_MULTIYEAR_WAV_006_012 time series (current BC: Hs=0.5 m, Tp=5 s, Dir=270° on the west).',
        'Resolve EDITO bucket quota and re-run the full pipeline on EDITO for cross-validation with the local v03a run.',
    ]:
        doc.add_paragraph(step, style='List Number')

    add_heading(doc, '7.3 Known risks', level=2)
    risks = [
        ('Hypersaline shock', 'Sharp 42→37.5 psu transition at inlets may trigger numerical overshoots; mitigation is polygon buffer widening if needed'),
        ('Evaporation double-counting', 'With Temperature=3 already computing latent flux, adding an explicit ERA5 evaporation forcing risks double-counting mass loss; verify convention in D-Flow FM 1.2.x release notes before committing'),
        ('EDITO parallel CPU', 'Pod default may be <4 CPUs; running nPart=4 MPI procs under CPU contention is slower than nPart=1'),
        ('EDITO credentials expiring mid-run', 'STS tokens typically expire in 24 h; long runs are unaffected (container has its own mount) but the post-processing JupyterLab may need renewal'),
        ('Storage quota', 'v03 ≥30-day map.nc could be 50–90 GB; within 50 GiB limit only with reduced output frequency or selective variable subset'),
    ]
    add_table(doc,
              header=['Risk', 'Mitigation / monitoring'],
              rows=risks,
              col_widths_cm=[6, 11])

    add_page_break(doc)


def write_references(doc):
    add_heading(doc, '8. References', level=1)

    add_para(doc, 'Site-specific literature (in reference/papers/):', bold=True)
    for ref in [
        'Ciraolo G. & De Marchis M. (2009). Stagnone di Marsala modeling. Journal of Coastal Research.',
        'Ciraolo G. et al. (2009). Particle tracking in aquatic vegetation meadows. [Referenced for Lagrangian methodology in seagrass environments.]',
        'Di Marca G. et al. (2009). Turbidity dynamics at Stagnone di Marsala.',
        'Emanuele et al. (2024). Water quality / hydrodynamic review (Water, 16, 2602).',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    add_para(doc, 'Methodology references:', bold=True)
    for ref in [
        'Deleersnijder E. & Delhez E.J.M. (2004). Residence time in semi-enclosed marine systems. JMS.',
        'Cucco A. & Umgiesser G. (2006). Modeling the Venice Lagoon residence time. JMS.',
        'Ferrarin C. et al. (2013). Sacca di Goro transport timescales.',
        'García-Oliva M. et al. (2019). Mar Menor hydrodynamic modeling under wind forcing. Ocean & Coastal Management.',
        'Liu Y. & Weisberg R.H. (2011). Evaluation of trajectory model predictions. Lagrangian skill score.',
        'Viero D.P. & Defina A. (2016). Water age in wind-driven lagoons: ensemble methodology. Advances in Water Resources.',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    add_para(doc, 'Technical / software references:', bold=True)
    for ref in [
        'Deltares (2026). Delft3D FM Suite 2026.01 HMWQ release notes. Deltares Open Source.',
        'Deltares. dfm_tools Python package — https://github.com/Deltares/dfm_tools',
        'Deltares. hydrolib-core 1.0.0 — MDU / ext / BC file parser.',
        'Delft3D User Days 2025 break-out session materials (in reference/).',
    ]:
        doc.add_paragraph(ref, style='List Bullet')

    add_page_break(doc)


def write_appendices(doc):
    add_heading(doc, 'Appendix A — Complete notebook catalog', level=1)

    # full table with purpose text
    rows = []
    for nb_file in sorted(NOTEBOOKS_DIR.glob('*.ipynb')):
        title, purpose, n_cells = read_notebook_first_markdown(nb_file)
        # truncate purpose
        if len(purpose) > 200:
            purpose = purpose[:197] + '...'
        rows.append([nb_file.name, title, purpose or '(no description)'])
    add_table(doc,
              header=['File', 'Title', 'Purpose (from first markdown cell)'],
              rows=rows,
              col_widths_cm=[5, 5, 7])

    add_page_break(doc)

    add_heading(doc, 'Appendix B — EDITO operational workflow', level=1)
    wf_path = PROJECT_ROOT / 'docs' / 'EDITO_WORKFLOW.md'
    if wf_path.exists():
        body = wf_path.read_text(encoding='utf-8')
        # Render in-line (skip top-level heading since we have one)
        in_code = False
        for line in body.splitlines():
            if line.startswith('```'):
                in_code = not in_code
                continue
            if line.startswith('# '):
                continue  # skip the top title, we have our own
            elif line.startswith('## '):
                add_heading(doc, line[3:].strip(), level=2)
            elif line.startswith('### '):
                add_heading(doc, line[4:].strip(), level=3)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif in_code:
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
            elif line.strip():
                add_para(doc, line)

    add_page_break(doc)

    add_heading(doc, 'Appendix C — Project memory (lessons learned)', level=1)
    memory_files = [
        ('project_stagnone_dt.md', 'Project scope and publications plan'),
        ('project_v01_calibration_notes.md', 'v01 calibration findings'),
        ('project_particle_tracking.md', 'July 2025 drifter campaign'),
        ('project_residence_time.md', 'Residence-time research question'),
        ('feedback_salinity_bc.md', 'Hypersaline salinity correction'),
        ('feedback_edito_run_model_sh.md', 'EDITO run_model.sh requirement'),
        ('reference_edito.md', 'EDITO Datalab account and quotas'),
    ]
    for fname, descr in memory_files:
        fm, body = read_memory(fname)
        if not body:
            continue
        add_heading(doc, f'{descr}', level=2)
        add_para(doc, f'Source: memory/{fname}', italic=True)
        # Render body plainly (strip markdown code fences)
        for block in body.split('\n\n'):
            block = block.strip()
            if not block:
                continue
            if block.startswith('**'):
                bold_text = block.strip('*').split('\n')[0]
                rest = '\n'.join(block.split('\n')[1:])
                add_para(doc, bold_text, bold=True)
                if rest:
                    add_para(doc, rest)
            else:
                add_para(doc, block)


# ---------- main ----------

def main():
    DOCS_DIR.mkdir(exist_ok=True)
    doc = Document()

    # default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    print('Generating cover...')
    write_cover(doc)

    print('Generating TOC placeholder...')
    write_toc_placeholder(doc)

    print('Section 1 — Executive summary...')
    write_executive_summary(doc)

    print('Section 2 — Introduction...')
    write_introduction(doc)

    print('Section 3 — Methodology...')
    write_methodology(doc)

    print('Section 4 — Implementation...')
    write_implementation(doc)

    print('Section 5 — Preliminary results...')
    write_results(doc)

    print('Section 6 — Research extensions...')
    write_extensions(doc)

    print('Section 7 — Current status + next steps...')
    write_status(doc)

    print('Section 8 — References...')
    write_references(doc)

    print('Appendices...')
    write_appendices(doc)

    doc.save(str(OUTPUT))
    print(f'\n✓ Report saved: {OUTPUT}')
    print(f'  Size: {OUTPUT.stat().st_size / 1024:.1f} KB')
    print('\nNext: open in Word, Insert → Table of Contents → Automatic Table 1, '
          'then polish manually.')


if __name__ == '__main__':
    main()
